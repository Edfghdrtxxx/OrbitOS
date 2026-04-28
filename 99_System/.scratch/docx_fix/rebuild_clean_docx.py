from pathlib import Path
import re
from docx import Document
from docx.document import Document as DocumentClass
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import Pt, Mm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

src = Path('D:/obsidian/OrbitOS/50_Resources/ComputerScience/扫描文稿 2026-4-28 21.48.13_副本.docx')
out = src.with_name(src.stem + '_cleaned_v3.docx')
if out.exists():
    raise SystemExit(f'输出文件已存在，避免覆盖: {out}')

CJK = '㐀-鿿豈-﫿'
SUP_SUB = str.maketrans({
    '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4', '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9',
    '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4', '₅': '5', '₆': '6', '₇': '7', '₈': '8', '₉': '9',
    '–': '-', '—': '-', '−': '-',
})

artifact_patterns = [
    re.compile(r'^(?:屏幕|解幕|幕)?\s*\d+\s*[-—]\s*\d+\s*[,，]?\s*共\s*\d+\s*个$'),
    re.compile(r'^\d{1,2}\s*[:：]\s*\d{2}(?:\s+\d{4}\s*/\s*\d{1,2}\s*/\s*\d{1,2})?$'),
    re.compile(r'^\d{4}\s*/\s*\d{1,2}\s*/\s*\d{1,2}$'),
    re.compile(r'^\+?\d{1,3}\s*%$'),
]

symbol_only = re.compile(r'^[□■◇◆●○★☆◎▫▭△▲▽▼·•\-—_\s]+$')


def iter_block_items(parent):
    if isinstance(parent, DocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent.element
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def collapse_spaced_acronyms(text: str) -> str:
    def repl(match):
        return re.sub(r'\s+', '', match.group(0))
    return re.sub(r'(?<![A-Za-z])(?:[A-Z]\s+){1,}[A-Z]{1,6}(?![a-z])', repl, text)


def clean_text(text: str, *, cell=False) -> str:
    if not text:
        return ''
    t = text.replace('\r', ' ').replace('\n', ' ')
    t = t.translate(SUP_SUB)
    t = t.replace('　', ' ')
    t = re.sub(r'[\t ]+', ' ', t)
    t = collapse_spaced_acronyms(t)
    t = re.sub(r'^\s*S\s+(?=(?:gs|gsql|nohup|source|chroot|sed|curl|ping)\b)', '$ ', t)
    t = re.sub(r'^\s*Sgs\b', '$ gs', t)
    t = re.sub(r'^\s*\$gs\b', '$ gs', t)
    t = re.sub(rf'(?<=[{CJK}])\s+(?=[{CJK}])', '', t)
    t = re.sub(rf'(?<=[{CJK}])\s+([，。；：、？！）\)】\]》])', r'\1', t)
    t = re.sub(rf'([（\(【\[《])\s+(?=[{CJK}])', r'\1', t)
    t = re.sub(r'\s+([,.;:?!\)])', r'\1', t)
    t = re.sub(r'([\(])\s+', r'\1', t)
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def compact(text: str) -> str:
    return re.sub(r'\s+', '', text)


def is_artifact(text: str, *, cell=False) -> bool:
    t = clean_text(text, cell=cell)
    if not t:
        return True
    c = compact(t)
    if not cell and c in {'I', 'l', '|', 'D', 'eno', 'm', '■', '□', '◇', '●', '★', '◎', 'Q搜索', '搜索', '英画品品一100%', '品十100%'}:
        return True
    if not cell and len(c) <= 2 and symbol_only.fullmatch(t):
        return True
    if not cell and any(p.fullmatch(t) or p.fullmatch(c) for p in artifact_patterns):
        return True
    if not cell and re.search(r'(屏幕|解幕|幕)\s*\d+\s*[-—]\s*\d+\s*[,，]?\s*共\s*\d+\s*个', t):
        return True
    if not cell and re.fullmatch(r'\d{3,}', c):
        return True
    if not cell and '搜索' in c and len(c) <= 10:
        return True
    if not cell and re.fullmatch(r'[+\-]?100%+', c):
        return True
    if not cell and '100%' in c and len(c) <= 12:
        return True
    if not cell and any(ch in c for ch in '□■◇◆★◎') and len(c) <= 30:
        return True
    if not cell and re.search(r'[\\x00-\\x7f]', c) and re.search(r'[一-鿿]', c) and len(c) <= 28:
        chinese = len(re.findall(r'[一-鿿]', c))
        latin_digit = len(re.findall(r'[A-Za-z0-9]', c))
        if chinese <= 4 and latin_digit >= 2 and not re.match(r'^\d+(\.\d+)*', c):
            return True
    return False


def is_body_start(text: str) -> bool:
    c = compact(clean_text(text))
    return bool(re.fullmatch(r'1项目概述', c) or re.fullmatch(r'1[\.、]?项目概述', c))


def heading_level(text: str):
    t = clean_text(text)
    if re.match(r'^\d+\s+\S+', t):
        return 1
    match = re.match(r'^(\d+(?:\.\d+){1,4})\s*\S+', t)
    if match:
        depth = match.group(1).count('.') + 1
        return min(depth, 3)
    return None


def looks_like_code(text: str) -> bool:
    t = text.strip()
    if re.match(r'^(\$|#)\s+', t):
        return True
    if re.match(r'^(ping|curl|chroot|source|sed|nohup|gsql|gs\s+guc|create\s+user|alter\s+user|SELECT\b|cat\s+)\b', t, re.I):
        return True
    if (';' in t or '|' in t or '>' in t) and re.search(r'\b(gs|gsql|sed|grep|CREATE|ALTER|SELECT|TABLE|INDEX)\b', t, re.I):
        return True
    return False


def set_run_font(run, font='Microsoft YaHei', size=None, bold=None, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_style_font(style, font='Microsoft YaHei', size=None, bold=None):
    style.font.name = font
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def add_clean_paragraph(doc, text):
    lvl = heading_level(text)
    if lvl == 1:
        p = doc.add_paragraph(style='Heading 1')
    elif lvl == 2:
        p = doc.add_paragraph(style='Heading 2')
    elif lvl == 3:
        p = doc.add_paragraph(style='Heading 3')
    elif looks_like_code(text):
        p = doc.add_paragraph(style='Code')
    else:
        p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if re.match(r'^\d+[、.)）]', text):
        p.paragraph_format.left_indent = Mm(4)
    run = p.add_run(text)
    if looks_like_code(text) and lvl is None:
        set_run_font(run, 'Consolas', 9)
    return p


def cell_lines(cell):
    lines = []
    for paragraph in cell.paragraphs:
        t = clean_text(paragraph.text, cell=True)
        if t and not is_artifact(t, cell=True):
            lines.append(t)
    if not lines:
        t = clean_text(cell.text, cell=True)
        if t and not is_artifact(t, cell=True):
            lines.append(t)
    dedup = []
    for line in lines:
        if not dedup or dedup[-1] != line:
            dedup.append(line)
    return dedup


def add_clean_table(out_doc, src_table):
    rows = []
    for row in src_table.rows:
        vals = []
        for cell in row.cells:
            vals.append('\n'.join(cell_lines(cell)))
        while vals and not vals[-1].strip():
            vals.pop()
        if any(v.strip() for v in vals):
            rows.append(vals)
    if not rows:
        return False
    max_cols = max(len(r) for r in rows)
    if max_cols < 2:
        for r in rows:
            txt = clean_text(' '.join(r))
            if txt and not is_artifact(txt):
                add_clean_paragraph(out_doc, txt)
        return True
    table = out_doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, vals in enumerate(rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ''
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            text = vals[c_idx] if c_idx < len(vals) else ''
            for i, line in enumerate(text.split('\n')):
                if i:
                    paragraph.add_run('\n')
                run = paragraph.add_run(line)
                set_run_font(run, 'Microsoft YaHei', 7 if max_cols >= 7 else 8, bold=(r_idx == 0))
    out_doc.add_paragraph()
    return True

src_doc = Document(src)
out_doc = Document()
section = out_doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(18)
section.bottom_margin = Mm(18)
section.left_margin = Mm(16)
section.right_margin = Mm(16)

styles = out_doc.styles
set_style_font(styles['Normal'], 'Microsoft YaHei', 10.5)
set_style_font(styles['Heading 1'], 'SimHei', 16, True)
set_style_font(styles['Heading 2'], 'SimHei', 13.5, True)
set_style_font(styles['Heading 3'], 'SimHei', 12, True)
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
try:
    code_style = styles.add_style('Code', 1)
except ValueError:
    code_style = styles['Code']
set_style_font(code_style, 'Consolas', 9)

started = False
paragraphs_added = 0
tables_added = 0
artifacts_skipped = 0
previous = None
consecutive_artifacts = 0

for block in iter_block_items(src_doc):
    if isinstance(block, Paragraph):
        txt = clean_text(block.text)
        if not started:
            if is_body_start(txt):
                started = True
            else:
                continue
        if is_artifact(txt):
            artifacts_skipped += 1
            consecutive_artifacts += 1
            continue
        if previous == txt and not heading_level(txt):
            artifacts_skipped += 1
            consecutive_artifacts += 1
            continue
        add_clean_paragraph(out_doc, txt)
        previous = txt
        paragraphs_added += 1
        consecutive_artifacts = 0
    elif isinstance(block, Table):
        if not started:
            continue
        if add_clean_table(out_doc, block):
            tables_added += 1
            previous = None
            consecutive_artifacts = 0

out_doc.save(out)
print(f'输出文件: {out}')
print(f'保留段落: {paragraphs_added}')
print(f'保留表格: {tables_added}')
print(f'跳过疑似页眉/图标残片: {artifacts_skipped}')
